import io
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# We use python-docx here (simpler for a server environment)
# Layout mirrors the docx-js version

def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

STEEL   = '2C3E50'
STEEL_D = '1A252F'
ACCENT  = 'E8B84B'
MUTED   = '5A5A5A'
BORDER  = 'D4CFC5'
BG      = 'F4F1EB'
BG_W    = 'EDE8DF'
WHITE   = 'FFFFFF'
BLACK   = '1A1A1A'


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_bottom_border_only(cell, color=BORDER):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tcBorders.append(b)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, size=10, color=BLACK, font='Arial', align=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*hex_to_rgb(color))
    run.font.name = font
    if align:
        para.alignment = align
    return run


def add_para(cell_or_doc, text, bold=False, size=10, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0):
    if hasattr(cell_or_doc, 'paragraphs') and hasattr(cell_or_doc, 'add_paragraph'):
        para = cell_or_doc.add_paragraph()
    else:
        para = cell_or_doc.paragraphs[0] if cell_or_doc.paragraphs else cell_or_doc.add_paragraph()
        # clear existing
        for p in cell_or_doc.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        para = cell_or_doc.paragraphs[0]
        para.clear()

    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    add_run(para, text, bold=bold, size=size, color=color)
    return para


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def generate_quote_docx(d):
    """Generate quote Word doc, return BytesIO buffer."""
    from docx import Document
    from docx.shared import Inches, Cm, Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page setup — A4, 20mm margins
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(1.4)
    section.bottom_margin = Cm(2)

    # Remove default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # ── HEADER BANNER ──
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = 'Table Grid'
    tbl.autofit = False

    cw_total = Cm(17)
    left_w   = Cm(10.5)
    right_w  = Cm(6.5)

    # Row 0: company name | quote ref
    r0 = tbl.rows[0]
    r0.height = Twips(900)

    c00 = r0.cells[0]
    c00.width = left_w
    set_cell_bg(c00, STEEL_D)
    remove_cell_borders(c00)
    set_cell_margins(c00, top=160, bottom=160, left=240, right=120)
    p = c00.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, 'AUTO GATES VIC', bold=True, size=24, color=WHITE)
    p2 = c00.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    add_run(p2, 'Custom Automatic Gates  ·  Phillip Island & Frankston', size=8.5, color='AAAAAA')

    c01 = r0.cells[1]
    c01.width = right_w
    set_cell_bg(c01, STEEL_D)
    remove_cell_borders(c01)
    set_cell_margins(c01, top=160, bottom=160, left=120, right=240)
    c01.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    add_run(p, 'QUOTE', bold=True, size=14, color=ACCENT)
    p2 = c01.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(2)
    add_run(p2, d['ref'], size=10, color=WHITE)
    p3 = c01.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_after = Pt(0)
    add_run(p3, d['date'], size=8.5, color='AAAAAA')

    # Row 1: accent stripe
    r1 = tbl.rows[1]
    r1.height = Twips(55)
    for ci in range(2):
        c = r1.cells[ci]
        set_cell_bg(c, ACCENT)
        remove_cell_borders(c)
        set_cell_margins(c, top=0, bottom=0, left=0, right=0)
        c.paragraphs[0].clear()

    # Row 2: contact | valid
    r2 = tbl.rows[2]
    r2.height = Twips(380)
    c20 = r2.cells[0]
    c20.width = left_w
    set_cell_bg(c20, BG)
    remove_cell_borders(c20)
    set_cell_margins(c20, top=100, bottom=100, left=240, right=120)
    p = c20.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, '0408 997 935   ·   james@autogatevic.com.au   ·   autogatesvic.netlify.app', size=8, color=MUTED)

    c21 = r2.cells[1]
    c21.width = right_w
    set_cell_bg(c21, BG)
    remove_cell_borders(c21)
    set_cell_margins(c21, top=100, bottom=100, left=120, right=240)
    p = c21.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    add_run(p, f'Valid for {d["valid_days"]} days from {d["date"]}', size=8, color=MUTED)

    doc.add_paragraph()

    # ── TWO-COL INFO ──
    info = doc.add_table(rows=1, cols=2)
    info.autofit = False
    half = Cm(8.25)

    def info_block(cell, items):
        remove_cell_borders(cell)
        set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
        first = True
        for label, value in items:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(0)
            add_run(p, label.upper() + '  ', bold=True, size=7.5, color=MUTED)
            add_run(p, value, bold=False, size=9.5, color=BLACK)

    info_left  = info.rows[0].cells[0]
    info_right = info.rows[0].cells[1]
    info_left.width  = half
    info_right.width = half

    info_block(info_left, [
        ('Prepared For', ''),
        ('Name',    d['name']),
        ('Phone',   d['phone']),
        ('Email',   d['email']),
        ('Address', d['address']),
    ])
    info_block(info_right, [
        ('Gate Specifications', ''),
        ('Gate Type',      d['gate_type']),
        ('Width × Height', f"{d['gate_width']} × {d['gate_height']}"),
        ('Infill',         d['infill']),
        ('Motor',          d['motor']),
        ('Driveway',       f"{d['driveway']} — {d['slope']}"),
        ('Electrical Run', d['electrical']),
        ('Access',         f"{d['remotes']} + {d['access']}"),
    ])

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)
    sp.paragraph_format.space_after  = Pt(2)

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_after = Pt(4)
    from docx.oxml import OxmlElement
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), BORDER)
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Section heading
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(2)
    h.paragraph_format.space_after  = Pt(4)
    add_run(h, 'QUOTE BREAKDOWN', bold=True, size=10, color=STEEL)

    # ── ITEMS TABLE ──
    col_d = Cm(13.2)
    col_a = Cm(3.8)

    items_tbl = doc.add_table(rows=1, cols=2)
    items_tbl.autofit = False

    # Header
    hr = items_tbl.rows[0]
    hc0, hc1 = hr.cells[0], hr.cells[1]
    hc0.width = col_d
    hc1.width = col_a
    set_cell_bg(hc0, STEEL)
    set_cell_bg(hc1, STEEL)
    remove_cell_borders(hc0)
    remove_cell_borders(hc1)
    set_cell_margins(hc0, top=100, bottom=100, left=160, right=80)
    set_cell_margins(hc1, top=100, bottom=100, left=80, right=160)
    p = hc0.paragraphs[0]
    add_run(p, 'Description', bold=True, size=9.5, color=WHITE)
    p = hc1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(p, 'Amount', bold=True, size=9.5, color=WHITE)

    def add_group_row(tbl, label, total_str):
        row = tbl.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = col_d
        c1.width = col_a
        set_cell_bg(c0, STEEL)
        set_cell_bg(c1, STEEL)
        remove_cell_borders(c0)
        remove_cell_borders(c1)
        set_cell_margins(c0, top=90, bottom=90, left=160, right=80)
        set_cell_margins(c1, top=90, bottom=90, left=80, right=160)
        p = c0.paragraphs[0]
        add_run(p, label, bold=True, size=9.5, color=WHITE)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, total_str, bold=True, size=9.5, color=ACCENT)

    def add_item_row(tbl, desc, amt):
        row = tbl.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = col_d
        c1.width = col_a
        set_cell_bg(c0, BG_W)
        set_cell_bg(c1, BG_W)
        set_bottom_border_only(c0)
        set_bottom_border_only(c1)
        set_cell_margins(c0, top=55, bottom=55, left=300, right=80)
        set_cell_margins(c1, top=55, bottom=55, left=80, right=160)
        p = c0.paragraphs[0]
        add_run(p, desc, size=9, color=MUTED)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, amt, size=9, color=MUTED)

    def add_subtotal_row(tbl, label, value):
        row = tbl.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = col_d
        c1.width = col_a
        set_cell_bg(c0, BG)
        set_cell_bg(c1, BG)
        set_bottom_border_only(c0)
        set_bottom_border_only(c1)
        set_cell_margins(c0, top=70, bottom=70, left=160, right=80)
        set_cell_margins(c1, top=70, bottom=70, left=80, right=160)
        p = c0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, label, size=9, color=MUTED)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, value, size=9, color=BLACK)

    def add_total_row(tbl, total_str):
        row = tbl.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = col_d
        c1.width = col_a
        set_cell_bg(c0, STEEL_D)
        set_cell_bg(c1, STEEL_D)
        remove_cell_borders(c0)
        remove_cell_borders(c1)
        set_cell_margins(c0, top=110, bottom=110, left=160, right=80)
        set_cell_margins(c1, top=110, bottom=110, left=80, right=160)
        p = c0.paragraphs[0]
        add_run(p, 'TOTAL  (inc GST)', bold=True, size=11, color=WHITE)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, total_str, bold=True, size=13, color=ACCENT)

    for group in d['groups']:
        add_group_row(items_tbl, group['label'], group['total'])
        for item in group['items']:
            add_item_row(items_tbl, item['desc'], item['amt'])

    add_subtotal_row(items_tbl, 'Subtotal (ex GST)', d['subtotal'])
    add_subtotal_row(items_tbl, 'GST (10%)',         d['gst'])
    add_total_row(items_tbl, d['total'])

    # ── SPACER + DIVIDER ──
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_before = Pt(6)
    sp2.paragraph_format.space_after  = Pt(2)
    div2 = doc.add_paragraph()
    div2.paragraph_format.space_after = Pt(4)
    pPr2 = div2._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    b2 = OxmlElement('w:bottom')
    b2.set(qn('w:val'), 'single')
    b2.set(qn('w:sz'), '4')
    b2.set(qn('w:color'), BORDER)
    pBdr2.append(b2)
    pPr2.append(pBdr2)

    # ── BOTTOM: SITE VISIT + ACCEPTANCE ──
    bot = doc.add_table(rows=1, cols=2)
    bot.autofit = False
    bl = bot.rows[0].cells[0]
    br = bot.rows[0].cells[1]
    bl.width = Cm(8.25)
    br.width = Cm(8.25)
    remove_cell_borders(bl)
    remove_cell_borders(br)
    set_cell_margins(bl, top=0, bottom=0, left=0, right=200)
    set_cell_margins(br, top=0, bottom=0, left=200, right=0)

    # Left col
    p = bl.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, 'PREFERRED SITE VISIT', bold=True, size=9, color=STEEL)
    p2 = bl.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    add_run(p2, f"{d['preferred_date']}  ·  {d['preferred_time']}", size=9.5, color=BLACK)
    p3 = bl.add_paragraph()
    p3.paragraph_format.space_after = Pt(2)
    add_run(p3, 'CUSTOMER NOTES', bold=True, size=9, color=STEEL)
    p4 = bl.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    add_run(p4, d['notes'] or 'None provided.', size=9.5, color=BLACK)

    # Right col — acceptance
    p = br.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    add_run(p, 'ACCEPTANCE', bold=True, size=9, color=STEEL)
    p2 = br.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    add_run(p2, 'By signing below you accept this quote and authorise Auto Gates Vic to proceed with supply and installation as described.', size=8.5, color=MUTED)

    for label in ['Signature:', 'Date:']:
        pl = br.add_paragraph()
        pl.paragraph_format.space_before = Pt(10)
        pl.paragraph_format.space_after  = Pt(2)
        add_run(pl, label, size=8.5, color=MUTED)
        pline = br.add_paragraph()
        pline.paragraph_format.space_after = Pt(2)
        pPrL = pline._p.get_or_add_pPr()
        pBdrL = OxmlElement('w:pBdr')
        bL = OxmlElement('w:bottom')
        bL.set(qn('w:val'), 'single')
        bL.set(qn('w:sz'), '4')
        bL.set(qn('w:color'), BORDER)
        pBdrL.append(bL)
        pPrL.append(pBdrL)

    # ── TERMS ──
    sp3 = doc.add_paragraph()
    sp3.paragraph_format.space_before = Pt(6)
    div3 = doc.add_paragraph()
    div3.paragraph_format.space_after = Pt(4)
    pPr3 = div3._p.get_or_add_pPr()
    pBdr3 = OxmlElement('w:pBdr')
    b3 = OxmlElement('w:bottom')
    b3.set(qn('w:val'), 'single'); b3.set(qn('w:sz'), '4'); b3.set(qn('w:color'), BORDER)
    pBdr3.append(b3); pPr3.append(pBdr3)

    th = doc.add_paragraph()
    th.paragraph_format.space_after = Pt(3)
    add_run(th, 'TERMS & CONDITIONS', bold=True, size=8.5, color=MUTED)

    terms = [
        '1.  This quote is valid for 30 days from the date of issue.',
        '2.  A 50% deposit is required upon acceptance. Balance payable on completion.',
        '3.  All pricing includes GST. Final invoice may vary if site conditions differ from those described.',
        '4.  Auto Gates Vic warrants all workmanship for 12 months. Motor manufacturer warranty applies separately.',
        '5.  Customer is responsible for obtaining any council permits required prior to installation.',
        '6.  Cancellation within 48 hours of scheduled installation may incur a call-out fee.',
    ]
    for t in terms:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_after = Pt(2)
        add_run(tp, t, size=7.5, color=MUTED)

    # ── FOOTER ──
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(4)
    pPrF = fp._p.get_or_add_pPr()
    pBdrF = OxmlElement('w:pBdr')
    bF = OxmlElement('w:top')
    bF.set(qn('w:val'), 'single'); bF.set(qn('w:sz'), '2'); bF.set(qn('w:color'), BORDER)
    pBdrF.append(bF); pPrF.append(pBdrF)
    add_run(fp, f"Auto Gates Vic  ·  0408 997 935  ·  james@autogatevic.com.au        Quote {d['ref']}", size=7.5, color=MUTED)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
